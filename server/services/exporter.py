import io
from typing import List
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models.qa_pair import QAPair


def export_to_docx(project_name: str, qa_pairs: List[QAPair]) -> io.BytesIO:
    doc = Document()

    title = doc.add_heading(f"{project_name} — Audit Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    total = len(qa_pairs)
    answered = sum(1 for qa in qa_pairs if qa.status == "answered")
    not_found = sum(1 for qa in qa_pairs if qa.status == "not_found")

    summary_para = doc.add_paragraph()
    summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = summary_para.add_run(
        f"Total Questions: {total}  |  Answered: {answered}  |  Not Found: {not_found}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")

    for qa in qa_pairs:
        q_heading = doc.add_heading(level=2)
        q_run = q_heading.add_run(f"Q{qa.question_number}: {qa.original_question}")
        q_run.font.size = Pt(13)

        answer_text = qa.ai_answer if qa.ai_answer else "Not answered yet."
        if qa.status == "not_found":
            answer_text = "Not found in references."

        answer_para = doc.add_paragraph()
        answer_label = answer_para.add_run("Answer: ")
        answer_label.bold = True
        answer_label.font.size = Pt(11)
        answer_content = answer_para.add_run(answer_text)
        answer_content.font.size = Pt(11)

        if qa.is_edited:
            edited_para = doc.add_paragraph()
            edited_run = edited_para.add_run("(Manually edited)")
            edited_run.font.size = Pt(9)
            edited_run.font.italic = True
            edited_run.font.color.rgb = RGBColor(150, 150, 150)

        citations = qa.citations or []
        if citations:
            citation_para = doc.add_paragraph()
            citation_label = citation_para.add_run("Citations: ")
            citation_label.bold = True
            citation_label.font.size = Pt(10)

            for citation in citations:
                if isinstance(citation, dict):
                    doc_name = citation.get("doc", "Unknown")
                    snippet = citation.get("snippet", "")
                    cite_text = f"[{doc_name}]"
                    if snippet:
                        cite_text += f" — \"{snippet[:100]}...\""
                else:
                    cite_text = f"[{citation}]"

                cite_run = citation_para.add_run(f"\n  • {cite_text}")
                cite_run.font.size = Pt(10)
                cite_run.font.color.rgb = RGBColor(0, 102, 204)

        evidence = qa.evidence_snippets or []
        if evidence:
            evidence_para = doc.add_paragraph()
            ev_label = evidence_para.add_run("Evidence Snippets: ")
            ev_label.bold = True
            ev_label.font.size = Pt(10)

            for ev in evidence:
                if isinstance(ev, dict):
                    ev_doc = ev.get("doc", "Unknown")
                    ev_text = ev.get("text", "")
                    ev_run = evidence_para.add_run(f"\n  📄 [{ev_doc}]: \"{ev_text}\"")
                else:
                    ev_run = evidence_para.add_run(f"\n  📄 {ev}")
                ev_run.font.size = Pt(9)
                ev_run.font.italic = True
                ev_run.font.color.rgb = RGBColor(80, 80, 80)

        doc.add_paragraph("─" * 60)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
